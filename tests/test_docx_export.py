from pathlib import Path

from docx import Document

from writing_agent.tools.export import export_docx, markdown_to_docx


def test_markdown_to_docx_handles_headings_lists_and_tables() -> None:
    document = Document()
    markdown_to_docx(
        """# Title

## Section

- bullet
1. numbered

| A | B |
|---|---|
| x | y |
""",
        document,
    )

    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Title" in text
    assert "Section" in text
    assert "bullet" in text
    assert document.tables[0].cell(1, 0).text == "x"


def test_export_docx_writes_formatted_document(tmp_path) -> None:
    path = export_docx(
        "# Plan\n\n## Section\n\nContent",
        output_dir=tmp_path,
        title="Delivery Plan",
        metadata={
            "document_type": "proposal",
            "topic": "Smart forestry",
            "audience": "reviewers",
            "model_name": "qwen3.6:35b",
            "rag_mode": "hybrid",
            "collection": "demo",
            "thread_id": "thread-1",
        },
    )

    assert Path(path).exists()
    document = Document(str(path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Delivery Plan" in text
    assert "生成说明" in text
    assert "qwen3.6:35b" in text
