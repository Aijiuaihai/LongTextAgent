from docx import Document

from writing_agent.tools.docx_template import render_docx_template


def test_docx_template_replaces_placeholders(tmp_path) -> None:
    template = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Title: {{title}}")
    document.add_paragraph("{{body}}")
    document.save(str(template))

    output, warnings = render_docx_template(
        template_path=template,
        output_path=tmp_path / "out.docx",
        markdown="# Body",
        values={"title": "Plan"},
    )

    assert output.exists()
    assert warnings == []
    rendered = Document(str(output))
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "Plan" in text
    assert "Body" in text

