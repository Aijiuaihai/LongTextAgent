from docx import Document

from writing_agent.tools.docx_preflight import validate_docx_template


def test_docx_template_preflight_passes_valid_template(tmp_path) -> None:
    template = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph(
        "{{title}} {{topic}} {{document_type}} {{audience}} {{generated_at}} "
        "{{body}} {{model_name}} {{rag_mode}} {{collection}} {{thread_id}}"
    )
    document.save(str(template))

    result = validate_docx_template(template)

    assert result.status == "pass"
    assert result.missing_placeholders == []


def test_docx_template_preflight_warns_when_body_missing(tmp_path) -> None:
    template = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("{{title}} {{topic}} {{document_type}} {{audience}} {{generated_at}}")
    document.save(str(template))

    result = validate_docx_template(template)

    assert result.status == "warning"
    assert "{{body}}" in result.missing_placeholders
    assert result.warnings

