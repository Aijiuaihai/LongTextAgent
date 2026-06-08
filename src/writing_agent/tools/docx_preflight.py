"""DOCX template preflight checks."""

import json
from pathlib import Path
from typing import Any, Literal

from pydantic import BaseModel, Field

REQUIRED_PLACEHOLDERS = [
    "{{title}}",
    "{{topic}}",
    "{{document_type}}",
    "{{audience}}",
    "{{generated_at}}",
    "{{body}}",
]
RECOMMENDED_PLACEHOLDERS = [
    "{{model_name}}",
    "{{rag_mode}}",
    "{{collection}}",
    "{{thread_id}}",
]
REQUIRED_STYLES = [
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Normal",
    "List Bullet",
    "List Number",
    "Table Grid",
]


class DocxTemplateInspection(BaseModel):
    """Raw template inspection details."""

    template_path: str
    exists: bool
    is_docx: bool
    can_open: bool
    text: str = ""
    placeholders: list[str] = Field(default_factory=list)
    styles: list[str] = Field(default_factory=list)
    error: str = ""


class DocxTemplatePreflightResult(BaseModel):
    """Template preflight result."""

    status: Literal["pass", "warning", "fail"]
    template_path: str
    missing_placeholders: list[str] = Field(default_factory=list)
    missing_recommended_placeholders: list[str] = Field(default_factory=list)
    missing_styles: list[str] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)
    recommendations: list[str] = Field(default_factory=list)
    error: str = ""


def _paragraph_texts(document: Any) -> list[str]:
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return texts


def _extract_placeholders(text: str) -> list[str]:
    import re

    return sorted(set(re.findall(r"\{\{[a-zA-Z0-9_]+\}\}", text)))


def inspect_docx_template(template_path: Path | str) -> DocxTemplateInspection:
    """Inspect a docx template for text, placeholders, and styles."""

    path = Path(template_path)
    if not path.exists():
        return DocxTemplateInspection(
            template_path=str(path),
            exists=False,
            is_docx=path.suffix.lower() == ".docx",
            can_open=False,
            error="Template file does not exist.",
        )
    if path.suffix.lower() != ".docx":
        return DocxTemplateInspection(
            template_path=str(path),
            exists=True,
            is_docx=False,
            can_open=False,
            error="Template must be a .docx file.",
        )
    try:
        from docx import Document

        document = Document(str(path))
    except Exception as exc:
        return DocxTemplateInspection(
            template_path=str(path),
            exists=True,
            is_docx=True,
            can_open=False,
            error=f"Could not open template: {exc}",
        )
    text = "\n".join(_paragraph_texts(document))
    return DocxTemplateInspection(
        template_path=str(path),
        exists=True,
        is_docx=True,
        can_open=True,
        text=text,
        placeholders=_extract_placeholders(text),
        styles=sorted({style.name for style in document.styles}),
    )


def validate_style_mapping(
    template_path: Path | str,
    style_mapping_path: Path | str | None = None,
) -> list[str]:
    """Return mapped styles that are missing from the template."""

    inspection = inspect_docx_template(template_path)
    if not inspection.can_open:
        return REQUIRED_STYLES
    if style_mapping_path and Path(style_mapping_path).exists():
        mapping = json.loads(Path(style_mapping_path).read_text(encoding="utf-8"))
        required = set(mapping.values())
    else:
        required = set(REQUIRED_STYLES)
    return sorted(style for style in required if style not in set(inspection.styles))


def validate_docx_template(
    template_path: Path | str,
    required_placeholders: list[str] | None = None,
    style_mapping_path: Path | str | None = None,
) -> DocxTemplatePreflightResult:
    """Validate a docx template before export."""

    required = required_placeholders or REQUIRED_PLACEHOLDERS
    inspection = inspect_docx_template(template_path)
    if not inspection.exists or not inspection.is_docx or not inspection.can_open:
        return DocxTemplatePreflightResult(
            status="fail",
            template_path=inspection.template_path,
            missing_placeholders=required,
            missing_styles=REQUIRED_STYLES,
            error=inspection.error,
            recommendations=["Provide a readable .docx template file."],
        )

    placeholders = set(inspection.placeholders)
    missing = [item for item in required if item not in placeholders]
    missing_recommended = [
        item for item in RECOMMENDED_PLACEHOLDERS if item not in placeholders
    ]
    missing_styles = validate_style_mapping(template_path, style_mapping_path)
    warnings: list[str] = []
    recommendations: list[str] = []
    if "{{body}}" not in placeholders:
        warnings.append("Template does not contain {{body}}; body will be appended.")
        recommendations.append("Add {{body}} where generated markdown should be inserted.")
    if "目录" in inspection.text and "更新域" not in inspection.text:
        warnings.append("Template mentions a table of contents but lacks an update-field note.")
        recommendations.append("Tell Word users to right-click the TOC and update fields.")
    if missing_recommended:
        warnings.append("Some recommended generation metadata placeholders are missing.")
    if missing_styles:
        warnings.append("Some mapped Word styles are missing from the template.")

    hard_missing = [item for item in missing if item != "{{body}}"]
    if hard_missing:
        status: Literal["pass", "warning", "fail"] = "fail"
    elif warnings:
        status = "warning"
    else:
        status = "pass"
    return DocxTemplatePreflightResult(
        status=status,
        template_path=inspection.template_path,
        missing_placeholders=missing,
        missing_recommended_placeholders=missing_recommended,
        missing_styles=missing_styles,
        warnings=warnings,
        recommendations=recommendations,
    )
