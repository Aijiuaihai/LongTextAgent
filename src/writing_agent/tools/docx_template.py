"""DOCX template placeholder support."""

import json
from pathlib import Path
from typing import Any

from writing_agent.tools.export import markdown_to_docx

DEFAULT_STYLE_MAPPING = {
    "markdown_h1": "Heading 1",
    "markdown_h2": "Heading 2",
    "markdown_h3": "Heading 3",
    "paragraph": "Normal",
    "bullet": "List Bullet",
    "numbered": "List Number",
    "table": "Table Grid",
}


def load_style_mapping(path: Path | str | None = None) -> dict[str, str]:
    """Load a style mapping JSON file or return defaults."""

    if path is None or not Path(path).exists():
        return DEFAULT_STYLE_MAPPING.copy()
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    return {**DEFAULT_STYLE_MAPPING, **data}


def _replace_text(paragraph: Any, values: dict[str, str]) -> bool:
    text = paragraph.text
    changed = False
    for key, value in values.items():
        placeholder = "{{" + key + "}}"
        if placeholder in text:
            text = text.replace(placeholder, value)
            changed = True
    if changed:
        paragraph.text = text
    return changed


def render_docx_template(
    *,
    template_path: Path | str,
    output_path: Path | str,
    markdown: str,
    values: dict[str, Any],
    style_mapping_path: Path | str | None = None,
) -> tuple[Path, list[str]]:
    """Render a docx template with placeholders and markdown body."""

    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError("Install python-docx to render docx templates.") from exc

    load_style_mapping(style_mapping_path)
    document = Document(str(template_path))
    string_values = {key: "" if value is None else str(value) for key, value in values.items()}
    warnings: list[str] = []
    body_inserted = False

    for paragraph in list(document.paragraphs):
        if "{{body}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{body}}", "")
            markdown_to_docx(markdown, document)
            body_inserted = True
        else:
            _replace_text(paragraph, string_values)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text(paragraph, string_values)

    if not body_inserted:
        warnings.append("Template does not contain {{body}}; appended body to the end.")
        markdown_to_docx(markdown, document)

    document.add_paragraph("目录可在 Microsoft Word 中右键更新域生成。")
    resolved_output = Path(output_path)
    resolved_output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(resolved_output))
    return resolved_output, warnings

