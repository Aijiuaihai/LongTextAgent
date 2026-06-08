"""Document export helpers."""

import re
from datetime import datetime
from pathlib import Path


def _slugify(value: str, max_length: int = 80) -> str:
    slug = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", value.strip().lower(), flags=re.UNICODE)
    slug = re.sub(r"-+", "-", slug).strip("-")
    return (slug or "document")[:max_length].strip("-") or "document"


def export_markdown(
    markdown: str,
    output_dir: Path | str = "./outputs",
    title: str = "document",
) -> Path:
    """Export markdown content to a timestamped file."""

    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    output_path = Path(output_dir) / f"{timestamp}-{_slugify(title)}.md"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(markdown, encoding="utf-8")
    return output_path


def export_docx(
    markdown: str,
    output_dir: Path | str = "./outputs",
    title: str = "document",
) -> Path:
    """Export a simple docx version of markdown content."""

    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError("Install python-docx to export .docx files.") from exc

    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    output_path = Path(output_dir) / f"{timestamp}-{_slugify(title)}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading(title, level=1)
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        else:
            document.add_paragraph(stripped)
    document.save(str(output_path))
    return output_path
